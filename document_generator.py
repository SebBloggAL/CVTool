# document_generator.py
import os
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import docx
import docx.oxml
import re
from datetime import datetime


def set_document_font(doc):
    """
    Sets the default font for the entire document to Calibri with a size of 10.5 pt.
    """
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')


def set_heading_style(doc):
    """
    Creates or modifies a style named 'Style 1' for headings with size 16 pt, amber color, bold, and underlined.
    """
    styles = doc.styles
    try:
        heading_style = styles['Style 1']
    except KeyError:
        heading_style = styles.add_style('Style 1', WD_STYLE_TYPE.PARAGRAPH)

    font = heading_style.font
    font.name = 'Calibri'
    font.size = Pt(16)
    font.color.rgb = RGBColor(226, 106, 35)  # Amber color
    font.bold = True
    font.underline = True  # Underline the headings
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Ensure paragraph formatting is consistent
    paragraph_format = heading_style.paragraph_format
    paragraph_format.space_after = Pt(12)


def set_document_defaults_language(doc):
    """
    Sets the document-wide default language to British English.
    """
    styles_element = doc.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    lang_default = rpr_default.xpath('w:lang')
    if lang_default:
        lang = lang_default[0]
    else:
        lang = docx.oxml.shared.OxmlElement('w:lang')
        rpr_default.append(lang)
    lang.set(qn('w:val'), 'en-GB')
    lang.set(qn('w:eastAsia'), 'en-US')
    lang.set(qn('w:bidi'), 'ar-SA')


def set_styles_language(doc):
    """
    Sets the language for all styles to British English.
    """
    for style in doc.styles:
        if style.type in [WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER]:
            rpr = style.element.get_or_add_rPr()
            lang = rpr.find(qn('w:lang'))
            if lang is None:
                lang = docx.oxml.shared.OxmlElement('w:lang')
                rpr.append(lang)
            lang.set(qn('w:val'), 'en-GB')
            lang.set(qn('w:eastAsia'), 'en-US')
            lang.set(qn('w:bidi'), 'ar-SA')


def apply_run_font_style(run, paragraph, is_applicant_name=False):
    """
    Applies font and language settings to a run.
    """
    if not run:
        return
    font = run.font
    font.name = 'Calibri'
    if paragraph.style.name == 'Style 1':
        font.size = Pt(16)
        font.underline = True
    elif is_applicant_name:
        font.size = Pt(20)
    else:
        font.size = Pt(10.5)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Set the language for the run using w:lang element
    rpr = run._element.get_or_add_rPr()
    lang = rpr.find(qn('w:lang'))
    if lang is None:
        lang = docx.oxml.shared.OxmlElement('w:lang')
        rpr.append(lang)
    lang.set(qn('w:val'), 'en-GB')
    lang.set(qn('w:eastAsia'), 'en-US')
    lang.set(qn('w:bidi'), 'ar-SA')


def replace_placeholders_in_paragraph(paragraph, placeholders):
    """
    Replaces placeholders in a paragraph with actual data.
    """
    full_text = ''.join(run.text for run in paragraph.runs)
    original_text = full_text

    for key, value in placeholders.items():
        if key in full_text:
            full_text = full_text.replace(key, value)
            logging.debug(f"Replaced '{key}' with '{value}' in paragraph.")

    if full_text != original_text:
        # Clear existing runs
        for run in paragraph.runs:
            p = run._element
            p.getparent().remove(p)
        paragraph._p.clear_content()
        paragraph.runs.clear()

        # Handle special formatting for {ApplicantName}
        if "{ApplicantName}" in original_text:
            parts = original_text.split("{ApplicantName}")
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                    apply_run_font_style(run, paragraph)
                if i < len(parts) - 1:
                    run = paragraph.add_run(placeholders["{ApplicantName}"])
                    apply_run_font_style(run, paragraph, is_applicant_name=True)
        else:
            new_run = paragraph.add_run(full_text)
            apply_run_font_style(new_run, paragraph)


def replace_headers(doc):
    """
    Replaces header placeholders (e.g., [Summary], [Certifications]) with actual header text and applies 'Style 1'.
    """
    header_placeholders = {
        "[Security Clearance]": "Security Clearance:",
        "[Summary]": "Summary",
        "[Skills]": "Skills",
        "[Experience]": "Experience",
        "[Education]": "Education",
        "[Certifications]": "Certifications"
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in header_placeholders:
            paragraph.text = header_placeholders[text]
            paragraph.clear()
            paragraph.style = 'Style 1'
            new_run = paragraph.add_run(header_placeholders[text])

            font = new_run.font
            font.name = 'Calibri'
            font.size = Pt(16)
            font.bold = True
            font.underline = True
            font.color.rgb = RGBColor(226, 106, 35)
            font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

            rpr = new_run._element.get_or_add_rPr()
            lang = rpr.find(qn('w:lang'))
            if lang is None:
                lang = docx.oxml.shared.OxmlElement('w:lang')
                rpr.append(lang)
            lang.set(qn('w:val'), 'en-GB')
            lang.set(qn('w:eastAsia'), 'en-US')
            lang.set(qn('w:bidi'), 'ar-SA')


def replace_placeholders_in_cell(cell, placeholders, data):
    """
    Replaces placeholders in a single cell and sets font and language for all runs.
    """
    for paragraph in cell.paragraphs:
        if '{Skills}' in paragraph.text:
            insert_skills_section(paragraph, data.get('Skills', []))
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        elif '{Experience}' in paragraph.text:
            insert_experience_section(paragraph, data.get('Experience', []))
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        elif '{Certifications}' in paragraph.text:
            insert_certifications_section(paragraph, data.get('Certifications', []))
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        else:
            replace_placeholders_in_paragraph(paragraph, placeholders)
            convert_lines_to_bullets(paragraph)

    for nested_table in cell.tables:
        replace_placeholders_in_table(nested_table, placeholders, data)


def replace_placeholders_in_table(table, placeholders, data):
    """
    Replaces placeholders in all cells of a table.
    """
    for row in table.rows:
        for cell in row.cells:
            replace_placeholders_in_cell(cell, placeholders, data)


def set_font_for_all_text(doc, placeholders):
    """
    Sets font and language for all runs in the document.
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            is_applicant_name = run.text == placeholders.get("{ApplicantName}", "")
            apply_run_font_style(run, paragraph, is_applicant_name=is_applicant_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        is_applicant_name = run.text == placeholders.get("{ApplicantName}", "")
                        apply_run_font_style(run, paragraph, is_applicant_name=is_applicant_name)


def set_list_bullet_style(doc):
    """
    Defines the 'List Bullet' style if it doesn't exist.
    """
    styles = doc.styles
    if 'List Bullet' not in styles:
        list_bullet_style = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        list_bullet_style.base_style = styles['List Paragraph']
        list_bullet_style.paragraph_format.left_indent = Pt(12)
        list_bullet_style.paragraph_format.first_line_indent = Pt(-12)
        list_bullet_style.paragraph_format.space_after = Pt(0)
        list_bullet_style.paragraph_format.space_before = Pt(0)


def insert_paragraph_after(paragraph, text='', style=None):
    """
    Inserts a new paragraph after the given paragraph.
    """
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_run = new_paragraph.add_run(text)
    if style is not None:
        new_paragraph.style = style
    return new_paragraph


def convert_lines_to_bullets(paragraph):
    """
    Converts multi‐line paragraphs into bullet‐styled paragraphs if lines start with bullet chars.
    """
    lines = paragraph.text.split('\n')
    if len(lines) > 1:
        original_style = paragraph.style
        paragraph.text = ''
        prev_para = paragraph
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('-') or line.startswith('•') or line.startswith('*'):
                text = line.lstrip('-•*').strip()
                style = 'List Bullet'
            else:
                text = line
                style = original_style
            new_para = insert_paragraph_after(prev_para, text=text, style=style)
            prev_para = new_para
    else:
        # Single line with a leading bullet
        txt = paragraph.text.strip()
        if txt.startswith('-') or txt.startswith('•') or txt.startswith('*'):
            paragraph.text = txt.lstrip('-•*').strip()
            paragraph.style = 'List Bullet'


def clean_duration_string(duration_str):
    """
    Cleans the duration string by removing backslashes and non-printable characters.
    """
    if duration_str:
        duration_str = duration_str.replace('\\', '')
        duration_str = ''.join(c for c in duration_str if 32 <= ord(c) <= 126)
    return duration_str.strip()


def identify_date_format(date_str):
    """
    Identifies a date format given a string like "Jan 2020" or "01/2021".
    """
    date_patterns = {
        '%d/%m/%Y': r'^\d{1,2}/\d{1,2}/\d{4}$',
        '%m/%Y': r'^\d{1,2}/\d{4}$',
        '%b %Y': r'^[A-Za-z]{3} \d{4}$',    # Jan 2020
        '%B %Y': r'^[A-Za-z]+ \d{4}$',      # January 2020
        '%Y': r'^\d{4}$'
    }
    for fmt, pattern in date_patterns.items():
        if re.match(pattern, date_str):
            return fmt
    return None


def parse_end_date(duration_str):
    """
    Parses the end date from a duration string. Returns datetime for sorting.
    """
    try:
        if not duration_str:
            return datetime.min

        duration_str = clean_duration_string(duration_str)
        present_terms = ["Present", "Current", "Now", "Ongoing"]
        parts = re.split(r'\s*[-–—]\s*', duration_str)
        if len(parts) == 2:
            end_str = parts[1].strip()
        elif len(parts) == 1:
            end_str = parts[0].strip()
        else:
            return datetime.min

        if any(term.lower() == end_str.lower() for term in present_terms):
            return datetime.now()

        date_fmt = identify_date_format(end_str)
        if date_fmt:
            return datetime.strptime(end_str, date_fmt)
        else:
            return datetime.min

    except Exception as e:
        logging.error(f"Error parsing duration '{duration_str}': {e}", exc_info=True)
        return datetime.min


def sort_experiences(experience_data):
    """
    Sorts a list of experience objects by their parsed end date, newest first.
    """
    for item in experience_data:
        duration = item.get("Duration", "")
        try:
            item['_end_date'] = parse_end_date(duration)
        except:
            item['_end_date'] = datetime.min

    sorted_list = sorted(experience_data, key=lambda x: x.get('_end_date', datetime.min), reverse=True)
    for item in sorted_list:
        item.pop('_end_date', None)

    return sorted_list


def insert_skills_section(paragraph, skills_data):
    """
    Inserts the skills section as bullet points below the given paragraph.
    """
    prev_para = paragraph
    if not skills_data:
        return

    for skill in skills_data:
        if not skill.strip():
            continue
        bullet_para = insert_paragraph_after(prev_para, skill.strip(), style='List Bullet')
        apply_run_font_style(bullet_para.runs[0], bullet_para)
        prev_para = bullet_para

    # Blank line for spacing
    insert_paragraph_after(prev_para, "")

    # Remove original placeholder
    p_el = paragraph._element
    p_el.getparent().remove(p_el)


def insert_experience_section(paragraph, experience_data):
    """
    Inserts the experience section below the given paragraph, sorted newest-to-oldest.
    """
    if not experience_data:
        p_el = paragraph._element
        p_el.getparent().remove(p_el)
        return

    experience_data = sort_experiences(experience_data)
    prev_para = paragraph

    for item in experience_data:
        position = item.get("Position", "")
        company = item.get("Company", "")
        duration = item.get("Duration", "")
        responsibilities = item.get("Responsibilities", [])

        if not position and not company:
            continue

        title_parts = [position]
        if company:
            title_parts.append(f"at {company}")
        if duration:
            title_parts.append(f"({duration})")
        title = ' '.join(title_parts)

        role_para = insert_paragraph_after(prev_para, "")
        role_para.style = 'Normal'
        role_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        role_run = role_para.add_run(title)
        role_run.bold = True
        apply_run_font_style(role_run, role_para)

        prev_para = role_para

        if isinstance(responsibilities, list):
            for resp in responsibilities:
                bullet_para = insert_paragraph_after(prev_para, resp.strip(), style='List Bullet')
                apply_run_font_style(bullet_para.runs[0], bullet_para)
                prev_para = bullet_para
        elif isinstance(responsibilities, str) and responsibilities.strip():
            bullet_para = insert_paragraph_after(prev_para, responsibilities.strip(), style='List Bullet')
            apply_run_font_style(bullet_para.runs[0], bullet_para)
            prev_para = bullet_para

        # Blank line for spacing
        prev_para = insert_paragraph_after(prev_para, "")

    # Remove original placeholder
    p_el = paragraph._element
    p_el.getparent().remove(p_el)


def insert_certifications_section(paragraph, cert_list):
    """
    Inserts the Certifications section as bullet points below the given paragraph.
    """
    prev_para = paragraph
    if not cert_list:
        return

    for cert in cert_list:
        if not cert.strip():
            continue
        bullet_para = insert_paragraph_after(prev_para, cert.strip(), style='List Bullet')
        apply_run_font_style(bullet_para.runs[0], bullet_para)
        prev_para = bullet_para

    # Blank line for spacing
    insert_paragraph_after(prev_para, "")

    # Remove original placeholder
    p_el = paragraph._element
    p_el.getparent().remove(p_el)


def create_document(data, output_path):
    """
    Creates a Word document using the template and fills it with the extracted data.

    Parameters:
      - data (dict): The data to populate in the document (must include "Certifications").
      - output_path (str): Where to save the final .docx.
    """
    template_path = 'Documents/Template/blank_template.docx'

    # Ensure output directory exists
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Load the template
    try:
        doc = Document(template_path)
    except Exception as e:
        logging.error(f"Failed to load template: {e}", exc_info=True)
        raise

    # Set fonts, heading styles, language, and list style
    set_document_font(doc)
    set_heading_style(doc)
    set_document_defaults_language(doc)
    set_styles_language(doc)
    set_list_bullet_style(doc)

    # Prepare placeholders (Skills/Experience/Certs handled separately)
    placeholders = {
        "{ApplicantName}": data.get("ApplicantName", ""),
        "{Role}": data.get("Role", ""),
        "{SecurityClearance}": data.get("SecurityClearance", "Not specified"),
        "{Summary}": data.get("Summary", ""),
        "{Education}": data.get("Education", "")
        # Note: {Skills}, {Experience}, {Certifications} are handled by custom insertion functions
    }

    logging.info("Starting placeholder replacement.")
    logging.debug(f"Skills: {data.get('Skills', [])}")
    logging.debug(f"Experience: {data.get('Experience', [])}")
    logging.debug(f"Certifications: {data.get('Certifications', [])}")

    # 1) Replace paragraph‐based placeholders and insert lists
    for paragraph in doc.paragraphs:
        if '{Skills}' in paragraph.text:
            insert_skills_section(paragraph, data.get('Skills', []))

        elif '{Experience}' in paragraph.text:
            insert_experience_section(paragraph, data.get('Experience', []))

        elif '{Certifications}' in paragraph.text:
            insert_certifications_section(paragraph, data.get('Certifications', []))

        else:
            replace_placeholders_in_paragraph(paragraph, placeholders)
            convert_lines_to_bullets(paragraph)

    # 2) Replace placeholders inside tables
    for table in doc.tables:
        replace_placeholders_in_table(table, placeholders, data)

    # 3) Replace header placeholders ([Summary], [Certifications], etc.)
    replace_headers(doc)

    logging.info("Placeholder replacement completed.")

    # 4) Ensure correct font and language for all runs
    set_font_for_all_text(doc, placeholders)

    # 5) Save the document
    try:
        doc.save(output_path)
        logging.info(f"Document saved to {output_path}")
    except Exception as e:
        logging.error(f"Failed to save document: {e}", exc_info=True)
        raise
