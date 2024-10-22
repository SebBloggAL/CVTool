# document_generator.py

import os
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import docx
import docx.oxml

def set_document_font(doc):
    """
    Sets the default font for the entire document to Calibri with a size of 10.5 pt.
    """
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

def set_heading_style(doc):
    """
    Creates or modifies a style named 'Style 1' for headings with size 16 pt and amber color.
    """
    styles = doc.styles
    try:
        heading_style = styles['Style 1']
    except KeyError:
        heading_style = styles.add_style('Style 1', WD_STYLE_TYPE.PARAGRAPH)

    font = heading_style.font
    font.name = 'Calibri'
    font.size = Pt(16)
    font.color.rgb = RGBColor(226, 106, 35)  # Amber color
    font.bold = True
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Ensure paragraph formatting is consistent
    paragraph_format = heading_style.paragraph_format
    paragraph_format.space_after = Pt(12)

def set_document_defaults_language(doc):
    """
    Sets the document-wide default language to British English.
    """
    styles_element = doc.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    lang_default = rpr_default.xpath('w:lang')
    if lang_default:
        lang = lang_default[0]
    else:
        lang = docx.oxml.shared.OxmlElement('w:lang')
        rpr_default.append(lang)
    lang.set(qn('w:val'), 'en-GB')
    lang.set(qn('w:eastAsia'), 'en-US')
    lang.set(qn('w:bidi'), 'ar-SA')

def set_styles_language(doc):
    """
    Sets the language for all styles to British English.
    """
    for style in doc.styles:
        if style.type in [WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER]:
            rpr = style.element.get_or_add_rPr()
            lang = rpr.find(qn('w:lang'))
            if lang is None:
                lang = docx.oxml.shared.OxmlElement('w:lang')
                rpr.append(lang)
            lang.set(qn('w:val'), 'en-GB')
            lang.set(qn('w:eastAsia'), 'en-US')
            lang.set(qn('w:bidi'), 'ar-SA')

def apply_run_font_style(run, paragraph, is_applicant_name=False):
    """
    Applies font and language settings to a run.
    """
    font = run.font
    font.name = 'Calibri'
    if paragraph.style.name == 'Style 1':
        font.size = Pt(16)
    elif is_applicant_name:
        font.size = Pt(20)
    else:
        font.size = Pt(10.5)
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Set the language for the run using w:lang element
    rpr = run._element.get_or_add_rPr()
    lang = rpr.find(qn('w:lang'))
    if lang is None:
        lang = docx.oxml.shared.OxmlElement('w:lang')
        rpr.append(lang)
    lang.set(qn('w:val'), 'en-GB')
    lang.set(qn('w:eastAsia'), 'en-US')
    lang.set(qn('w:bidi'), 'ar-SA')

def replace_placeholders_in_paragraph(paragraph, placeholders):
    """
    Replaces placeholders in a paragraph with actual data and adjusts bullet point alignment.
    """
    # Combine all run texts to handle placeholders split across runs
    full_text = ''.join(run.text for run in paragraph.runs)
    original_text = full_text  # Keep original for comparison

    for key, value in placeholders.items():
        if key in full_text:
            full_text = full_text.replace(key, value)
            logging.debug(f"Replaced '{key}' with '{value}' in paragraph.")

    if full_text != original_text:
        # Clear existing runs properly
        for run in paragraph.runs:
            p = run._element
            p.getparent().remove(p)
        paragraph._p.clear_content()
        paragraph.runs.clear()

        # Handle special formatting for specific placeholders
        if "{ApplicantName}" in original_text:
            # Split the text to isolate {ApplicantName}
            parts = original_text.split("{ApplicantName}")
            # Create runs for each part
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                    apply_run_font_style(run, paragraph)
                if i < len(parts) - 1:
                    # Insert ApplicantName with font size 20
                    run = paragraph.add_run(placeholders["{ApplicantName}"])
                    apply_run_font_style(run, paragraph, is_applicant_name=True)
        else:
            # Add a new run with the replaced text
            new_run = paragraph.add_run(full_text)
            apply_run_font_style(new_run, paragraph)

    # Adjust bullet point alignment
    if paragraph.text.strip().startswith('-'):
        paragraph.paragraph_format.left_indent = Pt(0)  # No indentation
        paragraph.paragraph_format.first_line_indent = Pt(-12)  # Adjust as needed
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def replace_headers(doc):
    """
    Replaces header placeholders (e.g., [Summary]) with actual header text and applies 'Style 1'.
    """
    header_placeholders = {
        "[Security Clearance]": "Security Clearance:",
        "[Summary]": "Summary",
        "[Skills]": "Skills",
        "[Experience]": "Experience",
        "[Education]": "Education"
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in header_placeholders:
            # Replace the placeholder with actual header text
            paragraph.text = header_placeholders[text]

            # Clear existing runs
            paragraph.clear()

            # Apply 'Style 1' to the paragraph
            paragraph.style = 'Style 1'

            # Add new run with header text
            new_run = paragraph.add_run(header_placeholders[text])

            # Set the font and language for the new run
            font = new_run.font
            font.name = 'Calibri'
            font.size = Pt(16)
            font.bold = True
            font.color.rgb = RGBColor(226, 106, 35)  # Amber color
            font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

            # Set the language
            rpr = new_run._element.get_or_add_rPr()
            lang = rpr.find(qn('w:lang'))
            if lang is None:
                lang = docx.oxml.shared.OxmlElement('w:lang')
                rpr.append(lang)
            lang.set(qn('w:val'), 'en-GB')
            lang.set(qn('w:eastAsia'), 'en-US')
            lang.set(qn('w:bidi'), 'ar-SA')

def replace_placeholders_in_cell(cell, placeholders):
    """
    Replaces placeholders in a single cell and sets font and language for all runs.
    """
    for paragraph in cell.paragraphs:
        replace_placeholders_in_paragraph(paragraph, placeholders)
        # Set font and language for each run in the paragraph
        for run in paragraph.runs:
            is_applicant_name = run.text == placeholders.get("{ApplicantName}", "")
            apply_run_font_style(run, paragraph, is_applicant_name=is_applicant_name)

    for nested_table in cell.tables:
        replace_placeholders_in_table(nested_table, placeholders)

def replace_placeholders_in_table(table, placeholders):
    """
    Replaces placeholders in all cells of a table.
    """
    for row in table.rows:
        for cell in row.cells:
            replace_placeholders_in_cell(cell, placeholders)

def set_font_for_all_text(doc, placeholders):
    """
    Sets the font and language for all text in the document to Calibri and British English.
    """
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            is_applicant_name = run.text == placeholders.get("{ApplicantName}", "")
            apply_run_font_style(run, paragraph, is_applicant_name=is_applicant_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        is_applicant_name = run.text == placeholders.get("{ApplicantName}", "")
                        apply_run_font_style(run, paragraph, is_applicant_name=is_applicant_name)

def create_document(data, output_path):
    """
    Creates a Word document using the template and fills it with the extracted data.

    Parameters:
    - data (dict): The data to populate in the document.
    - output_path (str): The full path where the document will be saved.
    """
    template_path = 'Documents/Template/blank_template.docx'

    # Ensure output directory exists
    output_directory = os.path.dirname(output_path)
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    # Load the template document
    try:
        doc = Document(template_path)
    except Exception as e:
        logging.error(f"Failed to load template: {e}")
        raise

    # Set the default font, heading style, and language
    set_document_font(doc)
    set_heading_style(doc)
    set_document_defaults_language(doc)  # Set document-wide default language
    set_styles_language(doc)  # Set language for all styles

    # Prepare placeholders
    placeholders = {
        "{ApplicantName}": data.get("ApplicantName", ""),
        "{Role}": data.get("Role", ""),
        "{SecurityClearance}": data.get("SecurityClearance", "Not specified"),
        "{Summary}": data.get("Summary", ""),
        "{Skills}": data.get("Skills", ""),
        "{Experience}": data.get("Experience", ""),
        "{Education}": data.get("Education", "")
    }

    logging.info("Starting placeholder replacement.")

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, placeholders)

    # Replace placeholders in tables
    for table in doc.tables:
        replace_placeholders_in_table(table, placeholders)

    # Replace header placeholders and apply 'Style 1'
    replace_headers(doc)

    logging.info("Placeholder replacement completed.")

    # Set font and language for all text after placeholder replacement
    set_font_for_all_text(doc, placeholders)

    # Save the filled-in document
    try:
        doc.save(output_path)
        logging.info(f"Document saved to {output_path}")
    except Exception as e:
        logging.error(f"Failed to save document: {e}")
        raise
