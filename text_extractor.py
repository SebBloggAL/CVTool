# text_extractor.py

import os
import re
import logging
import pdfminer
from pdfminer.high_level import extract_text as extract_pdf_text
import docx2txt

# Suppress pdfminer warnings
logging.getLogger('pdfminer').setLevel(logging.WARNING)

def extract_text(file_path):
    """
    Extracts and normalizes text from a file (PDF or DOCX).
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == '.pdf':
        logging.info("Extracting text from PDF file using pdfplumber.")
        raw = extract_text_from_pdf(file_path)
        return normalize_text(raw)

    elif ext == '.docx':
        logging.info("Extracting text from DOCX file.")
        try:
            raw = extract_text_from_docx_precise(file_path)  # table-aware
        except Exception as e:
            logging.warning(f"Precise DOCX extractor failed: {e}; falling back to docx2txt")
            raw = extract_text_from_docx(file_path)
        return normalize_text(raw)

    else:
        raise ValueError(f"Unsupported file extension: {ext}")


def extract_text_from_pdf(file_path):
    """
    Extracts text from a PDF file using pdfplumber for cleaner layout.
    """
    try:
        text_chunks = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ''
                text_chunks.append(page_text)
        return '\n'.join(text_chunks)
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        raise


def extract_text_from_docx(file_path):
    """
    Fallback: Extracts text from a DOCX file using docx2txt.
    """
    try:
        text = docx2txt.process(file_path)
        return text
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        raise


def extract_text_from_docx_precise(file_path):
    """
    Precise DOCX extraction that preserves paragraph order and also reads tables.
    """
    import docx
    d = docx.Document(file_path)
    lines = []

    # Paragraphs
    for p in d.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # Tables (row-major order)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        lines.append(t)

    return "\n".join(lines)


def normalize_text(text: str) -> str:
    """
    Normalize extracted text:
      1) Remove hyphenation at line breaks
      2) Keep single line breaks (so bullets/lines survive)
      3) Collapse 3+ blank lines into two
      4) Normalise Windows newlines
    """
    if not isinstance(text, str):
        return ""
    # 0) Normalise \r\n
    text = text.replace('\r\n', '\n')

    # 1) Remove hyphens at line ends
    text = re.sub(r'-\n', '', text)

    # 2) DO NOT merge single line-breaks; keep them

    # 3) Collapse 3+ consecutive blanks into two
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text
