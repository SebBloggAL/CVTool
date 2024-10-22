# document_generator.py

import os
import logging
from docx import Document

def replace_placeholders_in_paragraph(paragraph, placeholders):
    """
    Replaces placeholders in a paragraph with actual data.
    """
    inline = paragraph.runs
    # Combine all run texts to handle placeholders split across runs
    full_text = ''.join(run.text for run in inline)
    original_text = full_text  # Keep original for comparison

    for key, value in placeholders.items():
        if key in full_text:
            full_text = full_text.replace(key, value)
            logging.debug(f"Replaced '{key}' with '{value}' in paragraph.")

    if full_text != original_text:
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ''

        # Add a new run with the replaced text
        paragraph.add_run(full_text)

def replace_placeholders_in_table(table, placeholders):
    """
    Replaces placeholders in all cells of a table.
    """
    for row in table.rows:
        for cell in row.cells:
            replace_placeholders_in_cell(cell, placeholders)

def replace_placeholders_in_cell(cell, placeholders):
    """
    Replaces placeholders in a single cell.
    """
    for paragraph in cell.paragraphs:
        replace_placeholders_in_paragraph(paragraph, placeholders)
    for table in cell.tables:
        replace_placeholders_in_table(table, placeholders)

def create_document(data):
    """
    Creates a Word document using the template and fills it with the extracted data.
    """
    template_path = 'Documents/Template/blank_template.docx'
    output_directory = 'Documents/Processed'
    applicant_name = data.get('ApplicantName', 'output').replace(" ", "_")
    output_path = os.path.join(output_directory, f"{applicant_name}_CV.docx")

    # Ensure output directory exists
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    # Load the template document
    try:
        doc = Document(template_path)
    except Exception as e:
        logging.error(f"Failed to load template: {e}")
        raise

    # Replace placeholders in the document with data
    placeholders = {
        "{ApplicantName}": data.get("ApplicantName", ""),
        "{Role}": data.get("Role", ""),
        "{SecurityClearance}": data.get("SecurityClearance", "Not specified"),
        "{Summary}": data.get("Summary", ""),  # Summary may be empty
        "{Skills}": data.get("Skills", ""),
        "{Experience}": data.get("Experience", ""),
        "{Education}": data.get("Education", "")
    }

    logging.info("Starting placeholder replacement.")

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, placeholders)

    # Replace placeholders in tables
    for table in doc.tables:
        replace_placeholders_in_table(table, placeholders)

    logging.info("Placeholder replacement completed.")

    # Save the filled-in document
    try:
        doc.save(output_path)
        logging.info(f"Document saved to {output_path}")
    except Exception as e:
        logging.error(f"Failed to save document: {e}")
        raise
